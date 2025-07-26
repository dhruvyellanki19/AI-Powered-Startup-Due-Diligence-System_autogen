import os

from agents.founder import FounderAgent
from agents.market import MarketAnalystAgent
from agents.product import ProductReviewerAgent
from agents.competitor import CompetitorScoutAgent
from agents.tech import TechAuditorAgent
from agents.investment import InvestmentAdvisorAgent
from agents.risk import RiskOfficerAgent
from agents.synthesis import SynthesizerAgent
from agents.writer import ReportWriterAgent
from agents.vc import VCPartnerAgent

# Optional: for Word export
from docx import Document

# Read the startup description
with open("inputs/startup_description.txt") as f:
    startup_desc = f.read()

print("=== AI-Powered Startup Due Diligence System ===\n")

# Initialize agents
founder = FounderAgent()
market_analyst = MarketAnalystAgent()
product_reviewer = ProductReviewerAgent()
competitor_scout = CompetitorScoutAgent()
tech_auditor = TechAuditorAgent()
investment_advisor = InvestmentAdvisorAgent()
risk_officer = RiskOfficerAgent()
synthesizer = SynthesizerAgent()
report_writer = ReportWriterAgent()
vc_partner = VCPartnerAgent()

# 1. Founder
founder_answer = founder.run(
    question="What is your long-term vision for this startup?",
    startup_description=startup_desc
)
print("Founder response:\n", founder_answer, "\n" + "-"*60)

# 2. Market Analyst
market_report = market_analyst.run(
    question="Provide a detailed due diligence analysis of the market opportunity for this startup.",
    startup_description=startup_desc
)
print("Market Analyst Report:\n", market_report, "\n" + "-"*60)

# 3. Product Reviewer
product_report = product_reviewer.run(
    question="Critically evaluate the product's differentiation, user need, and scalability.",
    startup_description=startup_desc
)
print("Product Review:\n", product_report, "\n" + "-"*60)

# 4. Competitor Scout
competitor_report = competitor_scout.run(
    question="List major competitors, compare features, pricing, and user traction.",
    startup_description=startup_desc
)
print("Competitor Analysis:\n", competitor_report, "\n" + "-"*60)

# 5. Tech Auditor
tech_report = tech_auditor.run(
    question="Evaluate the tech stack, feasibility, security, scalability, and team capability.",
    startup_description=startup_desc
)
print("Tech Audit:\n", tech_report, "\n" + "-"*60)

# 6. Risk Officer
risk_report = risk_officer.run(
    question="Identify any significant business, legal, or technical risks.",
    startup_description=startup_desc
)
print("Risk Analysis:\n", risk_report, "\n" + "-"*60)

# 7. Investment Advisor
investment_report = investment_advisor.run(
    question="Would you recommend investing in this startup? Why or why not?",
    startup_description=startup_desc
)
print("Investment Advisor Opinion:\n", investment_report, "\n" + "-"*60)

# 8. Synthesizer
synthesis_report = synthesizer.run(
    question="Summarize all previous findings into a due diligence summary.",
    startup_description=startup_desc
)
print("Synthesis:\n", synthesis_report, "\n" + "-"*60)

# 9. Report Writer
final_report = report_writer.run(
    question="Write a professional due diligence report for this startup.",
    startup_description=startup_desc
)
print("Due Diligence Report:\n", final_report, "\n" + "-"*60)

# 10. VC Partner
vc_decision = vc_partner.run(
    question="Based on this due diligence report, should we invest? Give a GO/NO-GO and justify.",
    startup_description=startup_desc
)
print("VC Decision:\n", vc_decision, "\n" + "-"*60)

# === OPTIONAL: Export to Word Doc ===
doc = Document()
doc.add_heading('Startup Due Diligence Report', 0)
sections = [
    ('Founder response', founder_answer),
    ('Market Analyst Report', market_report),
    ('Product Review', product_report),
    ('Competitor Analysis', competitor_report),
    ('Tech Audit', tech_report),
    ('Risk Analysis', risk_report),
    ('Investment Advisor Opinion', investment_report),
    ('Synthesis', synthesis_report),
    ('Due Diligence Report', final_report),
    ('VC Decision', vc_decision)
]
for title, content in sections:
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
os.makedirs("outputs", exist_ok=True)
doc.save("outputs/startup_due_diligence_report.docx")
print("✅ Report saved as outputs/startup_due_diligence_report.docx")
